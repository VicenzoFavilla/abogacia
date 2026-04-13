"""
Servicio de editor de documentos con auto-guardado.

Flujo:
  1. Flutter solicita GET /documentos/{id}/contenido
  2. Backend convierte .docx → HTML con mammoth (o lee el HTML guardado)
  3. Flutter muestra el HTML en flutter_quill
  4. Cada 2s (debounce) Flutter hace PUT /documentos/{id}/contenido con el HTML actualizado
  5. Backend guarda el HTML y crea una Revision (auto-guardado)
  6. Cuando el usuario quiere exportar, Backend convierte HTML → .docx con python-docx
"""
import os
import io
import hashlib
from datetime import datetime
from pathlib import Path
from typing import Optional

import mammoth
from docx import Document as DocxDocument
from docx.shared import Pt
import pdfplumber
import PyPDF2


def docx_to_html(filepath: str) -> str:
    """Convierte un archivo .docx a HTML para edición en Flutter."""
    with open(filepath, "rb") as f:
        result = mammoth.convert_to_html(f)
    return result.value  # HTML limpio


def html_to_docx(html: str, output_path: str) -> str:
    """
    Convierte HTML editado de vuelta a .docx.
    Nota: mammoth solo hace docx→html. Para html→docx usamos python-docx
    con una conversión básica (para producción considerar pandoc o libreoffice).
    """
    from html.parser import HTMLParser

    class TextExtractor(HTMLParser):
        def __init__(self):
            super().__init__()
            self.text_parts = []
            self.current_tag = None

        def handle_starttag(self, tag, attrs):
            self.current_tag = tag

        def handle_data(self, data):
            if data.strip():
                self.text_parts.append((self.current_tag, data.strip()))

    parser = TextExtractor()
    parser.feed(html)

    doc = DocxDocument()
    for tag, text in parser.text_parts:
        if tag in ("h1",):
            p = doc.add_heading(text, level=1)
        elif tag in ("h2",):
            p = doc.add_heading(text, level=2)
        elif tag in ("h3",):
            p = doc.add_heading(text, level=3)
        else:
            doc.add_paragraph(text)

    doc.save(output_path)
    return output_path


def extract_text_from_pdf(filepath: str) -> str:
    """Extrae texto completo de un PDF con pdfplumber (mejor que PyPDF2 para texto)."""
    text_parts = []
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
    except Exception:
        # Fallback a PyPDF2
        try:
            with open(filepath, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        text_parts.append(text)
        except Exception:
            pass
    return "\n\n".join(text_parts)


def extract_text_from_docx(filepath: str) -> str:
    """Extrae texto plano de un .docx."""
    try:
        doc = DocxDocument(filepath)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    except Exception:
        return ""


def extract_text_from_txt(filepath: str) -> str:
    """Lee texto plano de un .txt."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    except Exception:
        return ""


def extract_text(filepath: str, tipo: str) -> str:
    """Router de extracción según tipo de archivo."""
    t = (tipo or "").lower()
    if t in ("pdf",) or filepath.lower().endswith(".pdf"):
        return extract_text_from_pdf(filepath)
    elif t in ("docx", "doc") or filepath.lower().endswith((".docx", ".doc")):
        return extract_text_from_docx(filepath)
    elif t == "txt" or filepath.lower().endswith(".txt"):
        return extract_text_from_txt(filepath)
    return ""


def get_docx_as_html(filepath: str) -> str:
    """Devuelve un .docx como HTML, o texto plano si falla."""
    ext = Path(filepath).suffix.lower()
    if ext in (".docx", ".doc"):
        return docx_to_html(filepath)
    elif ext == ".txt":
        text = extract_text_from_txt(filepath)
        # Envolver en párrafos HTML básicos
        paragraphs = [f"<p>{line}</p>" for line in text.split("\n") if line.strip()]
        return "\n".join(paragraphs)
    elif ext == ".pdf":
        text = extract_text_from_pdf(filepath)
        paragraphs = [f"<p>{line}</p>" for line in text.split("\n") if line.strip()]
        return (
            "<p><em>Los PDFs no son editables directamente. "
            "Este es el texto extraído solo lectura:</em></p>\n" + "\n".join(paragraphs)
        )
    return "<p>Formato de archivo no soportado para edición.</p>"


def compute_sha256(filepath: str) -> str:
    """Calcula el hash SHA-256 de un archivo para detección de duplicados."""
    sha256 = hashlib.sha256()
    with open(filepath, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            sha256.update(chunk)
    return sha256.hexdigest()
